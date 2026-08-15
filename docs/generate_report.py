# -*- coding: utf-8 -*-
"""Generates the LibMan SDLC project report as a .docx file."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x04, 0x16, 0x32)
GREY = RGBColor(0x44, 0x44, 0x44)

doc = Document()

# ---------- base styles ----------
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.15

for i, size in zip((1, 2, 3), (20, 15, 13)):
    st = doc.styles[f"Heading {i}"]
    st.font.name = "Calibri"
    st.font.size = Pt(size)
    st.font.color.rgb = NAVY
    st.font.bold = True
    st.paragraph_format.space_before = Pt(18 if i == 1 else 12)
    st.paragraph_format.space_after = Pt(8)

section = doc.sections[0]
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.top_margin = Inches(0.9)
section.bottom_margin = Inches(0.9)


def add_page_number_footer():
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld1)
    run._r.append(instr)
    run._r.append(fld2)


add_page_number_footer()


def h1(text, number=None):
    t = f"{number}. {text}" if number else text
    doc.add_heading(t, level=1)


def h2(text):
    doc.add_heading(text, level=2)


def h3(text):
    doc.add_heading(text, level=3)


def p(text, bold=False, italic=False, size=None, color=None, align=None):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if align:
        para.alignment = align
    return para


def bullets(items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def numbered(items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(htext)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "041632")
        hdr[i]._tc.get_or_add_tcPr().append(shading)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    doc.add_paragraph()
    return t


def page_break():
    doc.add_page_break()


# ============================================================
# TITLE PAGE
# ============================================================
for _ in range(6):
    doc.add_paragraph()

p("LibMan", bold=True, size=44, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)
p("University Library Management System", size=20, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
p("Project Report", bold=True, size=24, align=WD_ALIGN_PARAGRAPH.CENTER)
p("Prepared following the Software Development Life Cycle (SDLC)", italic=True, size=12,
  color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)

for _ in range(8):
    doc.add_paragraph()

p("Prepared by: Asante Gabriel Kwaku", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
p("Repository: github.com/AsanteAigu/LibMan", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
p("Date: August 2026", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)

page_break()

# ============================================================
# TABLE OF CONTENTS
# ============================================================
h1("Table of Contents")
toc_entries = [
    "1. Introduction",
    "2. System Requirements Analysis",
    "3. System Design",
    "4. Implementation",
    "5. Testing",
    "6. Deployment",
    "7. Maintenance and Future Enhancements",
    "8. Conclusion",
    "Appendix A: Database Schema Summary",
    "Appendix B: REST API Endpoint Reference",
    "Appendix C: Environment Configuration Reference",
]
for entry in toc_entries:
    doc.add_paragraph(entry, style="List Bullet")

page_break()

# ============================================================
# 1. INTRODUCTION
# ============================================================
h1("Introduction", 1)

h2("1.1 Project Overview")
p(
    "LibMan is a full-stack university Library Management System covering the complete "
    "lifecycle of library operations: catalogue browsing, physical book borrowing, digital "
    "(ebook) lending, reservations, fines and payments, and library staff administration. "
    "The system was built as a modern web application with a decoupled frontend and backend, "
    "a managed cloud database, real third-party payment processing, and a genuine cloud "
    "deployment across two hosting providers."
)

h2("1.2 Problem Statement")
p(
    "Traditional library management in many institutions still relies on manual record-keeping "
    "or legacy desktop software that does not support self-service borrowing, digital lending, "
    "or online fee payment. LibMan addresses this by providing a single web application that "
    "students can use to browse, borrow, read, and pay from any device, while giving librarians "
    "a real-time dashboard to manage the catalogue, approve requests, and track circulation."
)

h2("1.3 Objectives")
bullets([
    "Provide a self-service catalogue and borrowing workflow for students.",
    "Support both physical book circulation and digital ebook lending (PDF and EPUB) from the same platform.",
    "Automate fee and fine collection through a real payment gateway (Paystack), including a recurring monthly membership fee.",
    "Give librarians full catalogue, circulation, and user management tooling.",
    "Enforce business rules (hold windows, due dates, late fees, unpaid-charge blocking) consistently and automatically.",
    "Deploy the system to production infrastructure that remains continuously available on free-tier hosting.",
])

h2("1.4 Scope")
p(
    "The system supports two real user roles backed by the database (student and librarian; "
    "administrative functions are exposed to librarians rather than a separate role, since the "
    "schema does not distinguish an admin role), plus an unauthenticated guest view of the public "
    "catalogue. It covers the full borrow-request-to-return lifecycle for physical items, a "
    "parallel borrow-and-read lifecycle for ebooks, reservations for unavailable titles, a charges "
    "and payments subsystem, and a notifications system. It does not include acquisitions/ordering "
    "of new stock, inter-library loans, or a native mobile application (the system is instead built "
    "as an installable Progressive Web App)."
)

page_break()

# ============================================================
# 2. SYSTEM REQUIREMENTS ANALYSIS
# ============================================================
h1("System Requirements Analysis", 2)

h2("2.1 Requirements Gathering Approach")
p(
    "Requirements were captured iteratively against a design reference (a Google Stitch UI export) "
    "and an evolving understanding of the real data model. Where the original design brief exceeded "
    "what the database schema or design export supported, the missing screens were designed directly "
    "in the established visual language rather than left incomplete, and the decision to do so was "
    "confirmed with the project stakeholder before implementation."
)

h2("2.2 Functional Requirements")

h3("Guest / Public")
bullets([
    "Browse and search the full book catalogue without signing in.",
    "View a title's details, author, availability, and copy status.",
    "Register a new account or sign in to an existing one.",
])

h3("Student")
bullets([
    "Request to borrow an available copy, choosing a custom loan duration (1 minute to 30 days).",
    "Join a reservation queue for a title with no available copies.",
    "View active loans, due dates, and extend a loan once.",
    "Browse and borrow ebooks (PDF or EPUB) with a custom loan duration, and read them in-browser.",
    "View outstanding charges (late fees, damage/loss charges, membership fee) and pay them via Paystack.",
    "Receive notifications for reservation readiness, new charges, and expired holds.",
    "View and update a personal profile.",
]),

h3("Librarian")
bullets([
    "Approve or reject pending borrow requests.",
    "Mark a held copy as collected and process returns, including condition assessment.",
    "Add new titles and copies to the catalogue, including cover images.",
    "Upload ebook files (PDF/EPUB) and cover images for digital editions.",
    "Manage reservations and the librarian/user directory.",
    "Record cash payments on behalf of a student.",
    "View circulation reports and adjust system settings (e.g. the late fee rate).",
])

h2("2.3 Non-Functional Requirements")
table(
    ["Category", "Requirement"],
    [
        ("Availability", "Backend must remain reachable despite free-tier hosting's inactivity spin-down."),
        ("Security", "Passwords hashed (bcrypt); JWT-based stateless authentication; payment verification "
                      "performed server-side against Paystack, never trusting client-supplied success claims."),
        ("Usability", "Responsive layout for desktop and mobile; installable as a Progressive Web App."),
        ("Data integrity", "Core circulation rules (hold expiry, due dates, unpaid-charge blocking) enforced "
                            "at the database level via triggers, not solely in application code."),
        ("Portability", "Frontend and backend independently deployable to different cloud providers."),
        ("Maintainability", "Layered backend architecture (controller / service / repository); typed API "
                             "contracts on the frontend; automated typechecking and linting on every change."),
    ],
)

h2("2.4 User Roles and Permissions")
table(
    ["Role", "Key Permissions"],
    [
        ("Guest", "Browse catalogue and search; no borrowing or account-specific data."),
        ("Student (user)", "Borrow, reserve, read ebooks, pay charges; scoped strictly to their own data."),
        ("Librarian", "All student capabilities plus catalogue management, approvals, circulation "
                       "processing, user administration, and reporting/settings."),
    ],
)

page_break()

# ============================================================
# 3. SYSTEM DESIGN
# ============================================================
h1("System Design", 3)

h2("3.1 Architectural Overview")
p(
    "LibMan follows a decoupled client-server architecture. A Next.js single-page application "
    "communicates with a stateless Spring Boot REST API over HTTPS, authenticated with JSON Web "
    "Tokens. The API is the sole point of access to a managed PostgreSQL database hosted on "
    "Supabase, which also provides object storage for ebook files and cover images. Payments are "
    "processed through Paystack, with the backend performing independent server-side verification "
    "of every transaction before it is considered successful."
)
bullets([
    "Frontend: Next.js (App Router) single-page application, deployed as a Cloudflare Worker.",
    "Backend: Spring Boot REST API, deployed as a Docker container on Render.",
    "Database: PostgreSQL, hosted on Supabase (accessed via its session connection pooler).",
    "File storage: Supabase Storage (ebook files and cover images), accessed backend-only via a service-role credential.",
    "Payments: Paystack (Inline Popup checkout on the client, verified server-side).",
    "Continuous availability: a GitHub Actions scheduled workflow and an external uptime monitor "
    "both ping the backend's health endpoint to prevent Render's free-tier inactivity spin-down.",
])

h2("3.2 Technology Stack")
table(
    ["Layer", "Technology"],
    [
        ("Frontend framework", "Next.js 16.3 (App Router, Turbopack), React 19.2, TypeScript"),
        ("Frontend styling", "Tailwind CSS v4, shadcn/ui component primitives (built on Base UI)"),
        ("Frontend data/forms", "TanStack Query (server state), React Hook Form + Zod (forms/validation), Axios (HTTP)"),
        ("Ebook reading", "react-pdf (PDF rendering), epubjs (EPUB rendering)"),
        ("Backend framework", "Spring Boot 4.1, Spring Security 7.1 (OAuth2 Resource Server for JWT), Hibernate ORM 7.4"),
        ("Backend language/build", "Java 21, Maven"),
        ("Database", "PostgreSQL 17 (Supabase-hosted)"),
        ("File storage", "Supabase Storage"),
        ("Payments", "Paystack"),
        ("Backend hosting", "Render (Docker web service, free tier)"),
        ("Frontend hosting", "Cloudflare Workers (via the OpenNext Cloudflare adapter)"),
        ("CI / automation", "GitHub Actions (scheduled keep-alive workflow)"),
    ],
)

h2("3.3 Database Design")
p(
    "The schema is centred on titles and their physical copies, with parallel structures for "
    "ebook editions and ebook loans. Circulation state transitions are driven largely by "
    "PostgreSQL triggers, so that critical rules cannot be bypassed by a bug in application code."
)
h3("Core Tables")
bullets([
    "users -- account, role (librarian/user), hashed password.",
    "titles / copies -- catalogue entries and their individual physical copies.",
    "borrow_requests -- a student's request to borrow a copy, including a chosen loan duration.",
    "loans -- the physical hold-and-loan lifecycle: hold, collection, due date, return.",
    "reservations -- a per-title waiting queue for when no copy is available.",
    "ebook_editions / ebook_loans -- the digital equivalent, including uploaded file URL/format.",
    "charges / payments -- fines and fees, and the payments that clear them (cash or Paystack).",
    "notifications -- user-facing alerts (reservation ready, charge created, hold expired).",
    "settings -- librarian-configurable values, such as the late fee rate.",
])
h3("Key Database Triggers")
table(
    ["Trigger", "Responsibility"],
    [
        ("trg_check_unpaid_charges", "Blocks new borrow requests while any charge is unpaid."),
        ("trg_handle_borrow_request_approval", "Creates the loan (with a 6-hour collection window) on approval."),
        ("trg_handle_loan_collection", "Sets the due date from the requested duration once a copy is collected."),
        ("trg_handle_loan_return", "Releases the copy, or hands it to the next reservation in queue."),
        ("trg_auto_clear_charge", "Marks a charge paid the moment a payment is recorded against it."),
    ],
)

h2("3.4 API Design")
p(
    "The backend exposes a resource-oriented REST API under /api, secured with JWT bearer "
    "authentication. List endpoints are scoped to the authenticated user by default; a librarian "
    "token can request the unscoped view. Business-rule violations are returned as structured "
    "JSON errors ({\"error\": {\"code\", \"message\", \"fields\"}}), including translated "
    "PostgreSQL trigger exceptions, so the client never receives a raw stack trace."
)

h2("3.5 UI/UX Design Principles")
bullets([
    "A consistent typographic and colour system shared across all roles.",
    "Role-aware navigation: a public top navigation bar and a mobile bottom tab bar for patrons, "
    "a dedicated sidebar for librarian/admin screens.",
    "Progressive disclosure of complex actions behind confirmation dialogs (e.g. withdrawing a "
    "copy, rejecting a request) to prevent accidental destructive actions.",
    "Installable as a Progressive Web App with a dedicated icon, splash colours, and standalone "
    "display mode on both Android and iOS.",
])

page_break()

# ============================================================
# 4. IMPLEMENTATION
# ============================================================
h1("Implementation", 4)

h2("4.1 Development Methodology")
p(
    "The system was built iteratively, feature by feature, with each vertical slice (schema "
    "change, backend endpoint, frontend UI) implemented, compiled, and verified end-to-end before "
    "moving to the next. Given that several core frameworks (Spring Boot 4, Spring Security 7, "
    "Hibernate 7, Next.js 16) were newer than typical reference material, unfamiliar APIs were "
    "verified empirically -- by compiling frequently and inspecting IDE diagnostics -- rather than "
    "assumed from memory, which caught several real API-drift issues before they reached runtime."
)

h2("4.2 Frontend Implementation")
p(
    "The frontend is a single Next.js application with route groups separating public, "
    "authenticated-student, and librarian/admin areas, each behind its own layout and access "
    "guard. Server state (catalogue data, loans, charges, etc.) is managed exclusively through "
    "TanStack Query, giving automatic caching, invalidation, and background refresh. All forms use "
    "React Hook Form with Zod schemas for validation. A single typed Axios client attaches the "
    "JWT bearer token to every request and normalises backend error responses into a consistent "
    "shape the UI can render directly."
)

h2("4.3 Backend Implementation")
p(
    "The backend follows a layered architecture: REST controllers handle HTTP concerns only, "
    "service classes contain business logic, and Spring Data JPA repositories handle persistence. "
    "Authentication uses Spring Security's OAuth2 Resource Server module with a symmetric HS256 "
    "key, rather than a third-party JWT library, to avoid Jackson version-compatibility risk. "
    "Business rules are deliberately split between the database (via triggers, for rules that must "
    "hold regardless of which code path writes the data) and the application layer (for rules "
    "that need richer context, such as late-fee calculation and Paystack verification)."
)

h2("4.4 Core Feature Implementation")

h3("Catalogue and Borrowing")
p(
    "Students can request to borrow any title with an available copy, choosing their own loan "
    "duration at request time (from one minute up to thirty days). Once a librarian approves the "
    "request, a six-hour collection window opens; the due date is only set once the copy is "
    "physically marked collected, computed as collection time plus the requested duration. If a "
    "hold is never collected, a scheduled job automatically releases the copy back to the catalogue "
    "(or to the next student in the reservation queue) without penalising the student who never "
    "took the item."
)

h3("Ebook Lending and Reading")
p(
    "Librarians upload a PDF or EPUB file per ebook edition, stored in Supabase Storage via a "
    "backend-proxied upload (the browser never talks to Storage directly, since the application's "
    "own JWT authentication is separate from Supabase's own auth system). Students borrow an "
    "ebook with a self-chosen duration and read it directly in the browser: PDFs are rendered "
    "page-by-page with react-pdf inside the application's own styled container (rather than "
    "handing off to the browser's native PDF viewer), with defensive sizing so a single "
    "badly-scanned page cannot distort the layout of the rest of the book; EPUBs are rendered with "
    "epubjs. Both formats share a fullscreen reading mode."
)
p(
    "Reading access is actively enforced against the chosen duration, not just recorded for "
    "display. A scheduled job checks every minute for loans past their expiry, immediately moving "
    "them into a grace state and generating a payable charge; the reader page itself also runs a "
    "live client-side timer so access is cut the instant the duration elapses, rather than waiting "
    "for the next scheduled check. Paying the grace charge through the existing Paystack flow "
    "reactivates the loan for another full copy of the originally chosen duration; if the charge "
    "goes unpaid through a 24-hour grace window, the loan is removed outright and the edition "
    "becomes available for the student to borrow again from scratch."
)

h3("Charges, Fines, and Payments")
p(
    "Charges are created automatically: a late fee is calculated (days late times a "
    "librarian-configurable per-day rate) when an overdue item is returned; a replacement-cost "
    "charge is created for items returned damaged or lost; and a flat monthly membership fee is "
    "generated for every student, either by a scheduled job on the first of each month or lazily "
    "the moment a student attempts to borrow anything. Any unpaid charge blocks further borrowing, "
    "physical or digital, until it is cleared. Payment is made through a real Paystack Inline "
    "Popup checkout; the resulting transaction reference is never trusted on its own -- the "
    "backend independently re-verifies the transaction against Paystack's API, checking both that "
    "it succeeded and that the amount paid exactly matches the charge, before marking anything paid."
)

h3("Progressive Web App")
p(
    "The application is installable on both Android and iOS: a web manifest declares the app name, "
    "icons, and theme colours; Apple-specific meta tags enable a proper standalone launch "
    "experience on iOS; and a minimal service worker (scoped to same-origin requests only, so it "
    "never caches authenticated cross-origin API responses) provides the offline fallback needed "
    "for installability."
)

h2("4.5 Third-Party Integrations")
table(
    ["Service", "Purpose", "Integration Detail"],
    [
        ("Supabase", "Managed Postgres + file storage",
         "Postgres accessed via the session pooler; Storage accessed backend-only via a service-role key."),
        ("Paystack", "Payment processing",
         "Client-side Inline Popup checkout; server-side transaction verification before any charge is marked paid."),
        ("Cloudflare", "Frontend hosting",
         "Next.js deployed as a Cloudflare Worker via the OpenNext adapter."),
        ("Render", "Backend hosting",
         "Spring Boot deployed as a Docker container on Render's free web-service tier."),
        ("GitHub Actions", "Scheduled automation",
         "Pings the backend health endpoint every ten minutes to discourage inactivity spin-down."),
    ],
)

h2("4.6 Security Considerations")
bullets([
    "Passwords are hashed with bcrypt; no plaintext password is ever stored.",
    "Authentication is stateless JWT bearer tokens (HS256), scoped per request via Spring Security.",
    "Secrets (database credentials, JWT signing key, Supabase service-role key, Paystack secret "
    "key, Cloudflare API token) are kept exclusively in gitignored environment files or the "
    "hosting provider's secret store, never committed to source control.",
    "The Supabase service-role key and the Paystack secret key are used server-side only and are "
    "never exposed to the browser; only Paystack's public key (safe to expose by design) reaches "
    "client code.",
    "Payment success is never trusted from the client -- every payment is independently verified "
    "against Paystack's own API before a charge is cleared.",
    "CORS is restricted to an explicit allow-list of known frontend origins.",
])

page_break()

# ============================================================
# 5. TESTING
# ============================================================
h1("Testing", 5)

h2("5.1 Testing Strategy")
p(
    "Testing combined automated checks with deliberate, real end-to-end verification against the "
    "live Supabase database and, later, the deployed Paystack test environment and cloud hosting. "
    "Given that several business rules are enforced by database triggers rather than pure Java "
    "methods, they are not meaningfully unit-testable in isolation and were instead verified "
    "through real HTTP request walkthroughs (via curl) against the running application."
)

h2("5.2 Unit Testing")
p(
    "Pure business logic that does not depend on database triggers -- late-fee calculation and "
    "replacement-cost charging -- is covered by JUnit 5 and Mockito unit tests "
    "(LoanServiceTest). A Spring Boot context-load test confirms the application wires up "
    "correctly end to end against the real database configuration."
)

h2("5.3 Automated Static Verification")
p(
    "Every code change was validated with: TypeScript strict-mode type checking and ESLint on the "
    "frontend; a full Next.js production build (catching issues that only appear during static "
    "generation); and a Maven compile plus the JUnit test suite on the backend. These were run "
    "after every meaningful change throughout development, not only at the end."
)

h2("5.4 End-to-End / Manual Verification")
p(
    "Each major feature was exercised against the live system using real HTTP requests and, where "
    "relevant, real external services in test mode:"
)
bullets([
    "Authentication: registration, login, and password reset flows, including a real production "
    "bug where a directly-inserted plaintext password bypassed the bcrypt hashing path.",
    "Borrowing: physical borrow-request approval, collection, due-date calculation for several "
    "custom durations, extension, and return, including late-fee generation on a genuinely "
    "overdue return.",
    "Hold expiration: a real hold was backdated and the scheduled release job was run to confirm "
    "the copy correctly returned to availability (or to the next reservation).",
    "Ebook lending: file upload for both PDF and EPUB, borrowing with custom durations, and "
    "reading both formats in the browser.",
    "Membership fee and payments: a real Paystack test-mode transaction was created via Paystack's "
    "own API using a documented test card, then submitted to the backend and confirmed to verify, "
    "clear the charge, and unblock borrowing -- alongside confirming a fabricated payment "
    "reference is correctly rejected.",
    "Deployment connectivity: CORS, environment configuration, and the Paystack public key were "
    "each explicitly confirmed present and correct in the deployed production build, not just in "
    "local development.",
])

h2("5.5 Defects Identified and Resolved")
p("Selected defects found and fixed during development and testing, illustrative of the verification approach:")
table(
    ["Defect", "Root Cause", "Resolution"],
    [
        ("Search endpoints returned HTTP 500", "PostgreSQL could not infer a bind parameter's type "
         "from a bare comparison", "Added explicit CAST(... AS string) in the JPQL query"),
        ("Role filter silently matched zero rows", "Hibernate serialised a bare enum parameter by "
         "ordinal instead of name", "Passed the role as its String name instead of the enum type"),
        ("File uploads failed for every file", "Manually setting the multipart Content-Type header "
         "stripped the required boundary parameter", "Removed the explicit header; let the browser set it"),
        ("PDF reader showed huge blank gaps and a stretched black bar", "A single badly-scanned "
         "page's aspect ratio was reused as the size estimate for every other page in the book",
         "Each page now sizes independently, with a clamp against extreme aspect ratios"),
        ("Reader's fullscreen button was invisible", "The reader overlay had no z-index, so the "
         "site's own navigation bar painted over it", "Assigned the reader a z-index above all other UI layers"),
        ("A membership fee vanished if the same borrow attempt failed for an unrelated reason",
         "Fee creation shared a database transaction with the borrow attempt that could roll back",
         "Fee creation now commits in its own independent transaction"),
        ("Cloud build failed only on Linux, never locally", "An optional dependency's lockfile "
         "entry was incomplete because it was generated on Windows", "Explicitly pinned the missing "
         "transitive dependency so npm fully resolves it on every platform"),
    ],
)

page_break()

# ============================================================
# 6. DEPLOYMENT
# ============================================================
h1("Deployment", 6)

h2("6.1 Deployment Architecture")
p(
    "The frontend and backend are deployed independently to two different cloud providers, "
    "communicating over HTTPS with an explicit CORS allow-list. The database and file storage "
    "remain on Supabase throughout, so no data migration was required between deployment attempts."
)

h2("6.2 Backend Deployment (Render)")
p(
    "The backend is deployed as a Docker container on Render's free web-service tier. A "
    "multi-stage Dockerfile builds the Spring Boot application with a JDK image and runs it on a "
    "slim JRE image, giving full, explicit control over the Java runtime rather than depending on "
    "a hosting provider's auto-detected build environment. The database connection pool size is "
    "deliberately capped, since Supabase's free-tier session pooler allows only a small fixed "
    "number of total concurrent connections shared across every client on the project."
)

h2("6.3 Frontend Deployment (Cloudflare Workers)")
p(
    "The Next.js frontend is deployed as a Cloudflare Worker, built with the OpenNext Cloudflare "
    "adapter and deployed via Wrangler. This path was reached after two other hosting approaches "
    "were evaluated and set aside: a Vercel deployment (working, but ultimately consolidated onto "
    "a single provider) and Cloudflare's own Pages product using its built-in Next.js preset "
    "(@cloudflare/next-on-pages), which was found to have an unresolvable upstream dependency "
    "conflict with the current Wrangler release at the time of deployment."
)

h2("6.4 Environment Configuration and Secrets Management")
p(
    "Client-visible configuration (the API base URL, the Paystack public key) is provided as "
    "build-time environment variables, since Next.js inlines NEXT_PUBLIC_-prefixed variables into "
    "the compiled client bundle rather than reading them at request time. Server-only secrets "
    "(database credentials, the JWT signing secret, the Supabase service-role key, the Paystack "
    "secret key) are configured exclusively as environment variables on the hosting platform and "
    "are never present in source control."
)

h2("6.5 Continuous Availability (Keep-Alive Strategy)")
p(
    "Render's free tier spins a web service down after fifteen minutes without inbound HTTP "
    "traffic; an internal scheduled job within the application would not reset this timer, since "
    "only genuine external requests count as activity. Two independent, redundant mechanisms were "
    "put in place to prevent this:"
)
bullets([
    "A GitHub Actions workflow, scheduled to request the backend's health endpoint every ten "
    "minutes (noting that GitHub's own scheduler is best-effort and can be delayed, particularly "
    "on low-activity repositories).",
    "An external uptime monitor pinging the same health endpoint on a reliable five-minute "
    "interval, chosen specifically because it does not depend on GitHub's scheduling guarantees.",
])

page_break()

# ============================================================
# 7. MAINTENANCE AND FUTURE ENHANCEMENTS
# ============================================================
h1("Maintenance and Future Enhancements", 7)

h2("7.1 Current Limitations")
bullets([
    "Reservation fulfilment (when a copy becomes free for the next student in queue) does not "
    "currently offer that student a custom loan duration; it falls back to a fixed default.",
    "Deleting a user account directly at the database level bypasses the normal return workflow "
    "and can leave a copy's status inconsistent, since there is no in-application account-deletion "
    "feature to guard against this.",
    "The frontend's Cloudflare Worker deployment is currently triggered manually rather than "
    "automatically on every push, since Cloudflare's own Git-integrated build product could not "
    "be made to work reliably for this project at the time of deployment.",
    "The service worker provides installability and a basic offline shell, but does not "
    "pre-cache the full application for a rich offline experience.",
])

h2("7.2 Recommended Future Enhancements")
bullets([
    "Automated CI/CD for the frontend once Cloudflare's build tooling resolves its current "
    "dependency conflict, restoring push-to-deploy.",
    "Allow students to choose a loan duration when a reservation is fulfilled, matching the "
    "direct-borrow flow.",
    "Expand automated test coverage to include frontend component and integration tests.",
    "Add librarian-facing analytics for ebook readership and payment trends over time.",
    "Introduce push notifications (beyond in-app notifications) for due-date reminders, leveraging "
    "the PWA foundation already in place.",
])

page_break()

# ============================================================
# 8. CONCLUSION
# ============================================================
h1("Conclusion", 8)
p(
    "LibMan delivers a complete, production-deployed library management system spanning public "
    "catalogue browsing, physical and digital circulation, automated fee and fine handling with "
    "real payment processing, and a librarian administration suite, built on a modern decoupled "
    "architecture and verified through both automated checks and deliberate end-to-end testing "
    "against live infrastructure. The project also demonstrates the practical realities of "
    "shipping software on free-tier cloud infrastructure -- connection limits, inactivity "
    "spin-down, and cross-platform build inconsistencies -- and the engineering discipline of "
    "diagnosing each from first principles rather than guessing, before arriving at a working, "
    "continuously available deployment."
)

page_break()

# ============================================================
# APPENDIX A
# ============================================================
h1("Appendix A: Database Schema Summary")
table(
    ["Table", "Purpose"],
    [
        ("users", "Accounts: name, email, hashed password, role."),
        ("titles / copies", "Catalogue entries and their individual physical copies."),
        ("borrow_requests", "Student requests to borrow a copy, including chosen duration."),
        ("loans", "Physical hold/collection/due-date/return lifecycle."),
        ("reservations", "Per-title waiting queue."),
        ("ebook_editions", "Digital edition per title: uploaded file URL and format."),
        ("ebook_loans", "Digital borrow lifecycle: borrowed/expires/returned."),
        ("charges", "Fines and fees: late fee, damage, lost, membership fee."),
        ("payments", "Payments clearing a charge, by cash or Paystack."),
        ("notifications", "User-facing alerts."),
        ("settings", "Librarian-configurable values (e.g. late fee rate)."),
        ("withdrawal_log", "Audit trail of copies withdrawn from circulation."),
    ],
)

page_break()

# ============================================================
# APPENDIX B
# ============================================================
h1("Appendix B: REST API Endpoint Reference")
table(
    ["Endpoint", "Purpose"],
    [
        ("POST /api/auth/register, /login", "Account creation and authentication."),
        ("GET /api/titles, /api/titles/:id", "Catalogue browsing and search."),
        ("POST /api/titles, /api/titles/:id/copies", "Librarian catalogue management."),
        ("POST /api/titles/:id/cover", "Librarian cover image upload."),
        ("GET/POST /api/borrow-requests", "Borrow request submission and listing."),
        ("PATCH /api/borrow-requests/:id/approve|reject", "Librarian request decisions."),
        ("GET /api/loans", "Loan listing (self-scoped, or all for a librarian)."),
        ("PATCH /api/loans/:id/collect|return|extend", "Circulation processing."),
        ("GET/POST /api/reservations", "Reservation queue management."),
        ("GET /api/ebook-editions", "Ebook catalogue listing."),
        ("POST /api/ebook-editions/:id/file", "Librarian ebook file upload."),
        ("GET/POST /api/ebook-loans", "Ebook borrowing and listing."),
        ("GET /api/charges, /api/payments", "Charge and payment history."),
        ("POST /api/payments", "Submit a payment (cash or verified Paystack transaction)."),
        ("GET /api/notifications", "User notifications."),
        ("GET /api/users, /api/users/:id/history", "Librarian user directory and lookup."),
        ("GET/PUT /api/settings", "Librarian-configurable system settings."),
        ("GET /actuator/health", "Health check endpoint used by the keep-alive mechanisms."),
    ],
)

page_break()

# ============================================================
# APPENDIX C
# ============================================================
h1("Appendix C: Environment Configuration Reference")
h2("Backend")
table(
    ["Variable", "Purpose"],
    [
        ("DATABASE_URL / _USERNAME / _PASSWORD", "PostgreSQL connection (Supabase session pooler)."),
        ("JWT_SECRET / JWT_EXPIRATION_MS", "JWT signing key and token lifetime."),
        ("CORS_ALLOWED_ORIGINS", "Comma-separated list of allowed frontend origins."),
        ("SUPABASE_STORAGE_URL / SUPABASE_SERVICE_ROLE_KEY", "Backend-only file storage access."),
        ("PAYSTACK_SECRET_KEY", "Server-side payment verification."),
        ("PORT", "HTTP port, provided by the hosting platform at runtime."),
    ],
)
h2("Frontend")
table(
    ["Variable", "Purpose"],
    [
        ("NEXT_PUBLIC_API_BASE_URL", "Backend API base URL, baked in at build time."),
        ("NEXT_PUBLIC_PAYSTACK_PUBLIC_KEY", "Paystack checkout key; safe to expose client-side by design."),
    ],
)

doc.save(r"c:\Users\asant\Documents\LibMan\docs\LibMan_Project_Report.docx")
print("Report generated successfully.")
